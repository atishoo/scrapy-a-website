import scrapy
from docx import Document
from docx.shared import Cm
import hashlib

class ImageItem(scrapy.Item):
    image_urls = scrapy.Field()
    images = scrapy.Field()

class KaibogroupSpider(scrapy.Spider):
    name = 'kaibogroup'
    allowed_domains = []
    start_urls = ['http://www.kaibogroup.com/newscn.aspx?type=2']

    def parse(self, response):
        ''' 获取文章链接 '''
        list = response.css('#newslb>li')
        for listItem in list:
            articleUrl = listItem.css('a::attr(href)').get()
            if articleUrl is not None:
                yield response.follow(articleUrl, callback=self.parseArticle)

        ''' 获取分页链接 '''
        nextpage = response.css('#fy2 > a')[2].attrib['href']
        if (nextpage != response.css('#fy2 > a')[3].attrib['href']):
            yield response.follow(nextpage, callback=self.parse)


    def parseArticle(self, response):
        title = response.css('#ej2 h2::text').get()
        extraInfo = response.css('#ej2 h3::text').get()
        extraInfo = extraInfo.replace(' ', '')
        extraInfo = extraInfo.split("\r\n")
        author = extraInfo[1].replace('来自：', '')
        publishDate = extraInfo[3].replace('加入时间：', '')
        article = response.css('#ej2 td>div>*')
        ''' 下载图片 '''
        imgs = response.css('#ej2 img::attr(src)').getall()
        if len(imgs) > 0:
            item = ImageItem()
            item['image_urls'] = imgs
            yield item

        self.saveAsDocx(title, author, publishDate.replace('/', '-'), article, response.url)

    def saveAsDocx(self, title, author, publishDate, article, link):
        document = Document()

        document.add_heading(title, 0)
        document.add_paragraph(f'作者：{author}       发布时间：{publishDate}')
        document.add_paragraph('')

        ''' 解析文章内容 '''
        if len(article) > 0:
            for p in article:
                if len(p.css('#prev')) > 0:
                    continue
                elif p.css('img::attr(src)').get():
                    imgLink = p.css('img::attr(src)').get()
                    try:
                        suffix = imgLink.split('.')[-1]
                        filename = hashlib.sha1(imgLink.encode('utf-8')).hexdigest() + '.' + suffix
                        document.add_picture(f'./imgs/full/{filename}', width=Cm(16))
                    except Exception as e:
                        print('error: ', e)
                        print('asdfadsf::' + imgLink)
                else:
                    paragraph = document.add_paragraph(p.css('*::text').getall())

        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('摘自：' + link)

        prefix = publishDate.replace('-', '.')
        author = author.replace('/', '&')
        document.save(f'doc/{prefix}-{title}（{author}）.docx')

